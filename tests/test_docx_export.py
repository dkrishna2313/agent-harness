"""Tests for functional_agents.docx_export — markdown → DOCX conversion."""
from __future__ import annotations

import textwrap
from pathlib import Path

import pytest

from functional_agents.docx_export import convert, _add_runs, _page_break_before

pytest.importorskip("docx", reason="python-docx not installed")


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

_SAMPLE_MD = textwrap.dedent("""\
    # Executive Report

    ## 1. Executive Summary

    > **Recommendation:** Option A
    >
    > **Confidence:** High

    This is a regular paragraph with **bold** and *italic* text.

    ---

    ## 2. Strategic Context

    *What decision needs to be made?*

    | Field | Value |
    |---|---|
    | Option | Alpha |
    | Status | Active |

    - First bullet
    - Second bullet

    1. First item
    2. Second item

    ## 10. Appendix

    ### Appendix A — Detail

    #### Sub-section heading

    Appendix content.
""")


@pytest.fixture()
def sample_docx(tmp_path: Path) -> Path:
    md = tmp_path / "report.md"
    md.write_text(_SAMPLE_MD, encoding="utf-8")
    out = tmp_path / "report.docx"
    convert(md, out)
    return out


# ---------------------------------------------------------------------------
# File-level checks
# ---------------------------------------------------------------------------

def test_output_file_created(sample_docx):
    assert sample_docx.exists()
    assert sample_docx.stat().st_size > 0


def test_output_extension_docx(sample_docx):
    assert sample_docx.suffix == ".docx"


def test_default_output_path(tmp_path):
    md = tmp_path / "report.md"
    md.write_text(_SAMPLE_MD, encoding="utf-8")
    convert(md, md.with_suffix(".docx"))
    assert (tmp_path / "report.docx").exists()


# ---------------------------------------------------------------------------
# Content checks via python-docx
# ---------------------------------------------------------------------------

def _open(path: Path):
    from docx import Document
    return Document(str(path))


def test_heading_h1_present(sample_docx):
    doc = _open(sample_docx)
    h1_texts = [p.text for p in doc.paragraphs if p.style.name == "Title"]
    assert any("Executive Report" in t for t in h1_texts)


def test_heading_h2_present(sample_docx):
    doc = _open(sample_docx)
    h2_texts = [p.text for p in doc.paragraphs if p.style.name == "Heading 1"]
    assert any("Executive Summary" in t for t in h2_texts)


def test_heading_h3_present(sample_docx):
    doc = _open(sample_docx)
    h3_texts = [p.text for p in doc.paragraphs if p.style.name == "Heading 2"]
    assert any("Appendix A" in t for t in h3_texts)


def test_heading_h4_present(sample_docx):
    doc = _open(sample_docx)
    h4_texts = [p.text for p in doc.paragraphs if p.style.name == "Heading 3"]
    assert any("Sub-section" in t for t in h4_texts)


def test_table_present(sample_docx):
    doc = _open(sample_docx)
    assert len(doc.tables) >= 1


def test_table_header_row(sample_docx):
    doc = _open(sample_docx)
    table = doc.tables[0]
    header_texts = [cell.text for cell in table.rows[0].cells]
    assert "Field" in header_texts
    assert "Value" in header_texts


def test_table_body_row(sample_docx):
    doc = _open(sample_docx)
    table = doc.tables[0]
    assert table.rows[1].cells[0].text == "Option"
    assert table.rows[1].cells[1].text == "Alpha"


def test_paragraph_text_present(sample_docx):
    doc = _open(sample_docx)
    all_text = "\n".join(p.text for p in doc.paragraphs)
    assert "regular paragraph" in all_text


def test_blockquote_text_present(sample_docx):
    doc = _open(sample_docx)
    all_text = "\n".join(p.text for p in doc.paragraphs)
    assert "Option A" in all_text
    assert "High" in all_text


def test_bullet_list_present(sample_docx):
    doc = _open(sample_docx)
    all_text = "\n".join(p.text for p in doc.paragraphs)
    assert "First bullet" in all_text
    assert "Second bullet" in all_text


def test_ordered_list_present(sample_docx):
    doc = _open(sample_docx)
    all_text = "\n".join(p.text for p in doc.paragraphs)
    assert "First item" in all_text
    assert "Second item" in all_text


def test_no_raw_markdown_artifacts(sample_docx):
    doc = _open(sample_docx)
    all_text = "\n".join(p.text for p in doc.paragraphs)
    # No heading markers
    assert "## " not in all_text
    assert "# " not in all_text
    # No raw bold/italic markers in plain text
    assert "**" not in all_text
    # No table separator lines
    assert "|---|" not in all_text


# ---------------------------------------------------------------------------
# Unit tests for helpers
# ---------------------------------------------------------------------------

def test_page_break_before_appendix():
    assert _page_break_before("10. Appendix") is True
    assert _page_break_before("Appendix A — Detail") is True


def test_no_page_break_before_regular_section():
    assert _page_break_before("1. Executive Summary") is False
    assert _page_break_before("2. Strategic Context") is False


def test_add_runs_bold(tmp_path):
    from docx import Document
    doc = Document()
    p = doc.add_paragraph()
    _add_runs(p, "plain **bold text** more")
    texts = [r.text for r in p.runs]
    bolds = [r.bold for r in p.runs]
    assert "bold text" in texts
    assert bolds[texts.index("bold text")] is True


def test_add_runs_italic(tmp_path):
    from docx import Document
    doc = Document()
    p = doc.add_paragraph()
    _add_runs(p, "plain *italic text* more")
    texts = [r.text for r in p.runs]
    italics = [r.italic for r in p.runs]
    assert "italic text" in texts
    assert italics[texts.index("italic text")] is True


def test_add_runs_plain(tmp_path):
    from docx import Document
    doc = Document()
    p = doc.add_paragraph()
    _add_runs(p, "just plain text")
    assert p.text == "just plain text"


# ---------------------------------------------------------------------------
# Real report smoke test (skipped if file absent)
# ---------------------------------------------------------------------------

_REAL_REPORT = Path("outputs/sports_strategy.md")


@pytest.mark.skipif(
    not _REAL_REPORT.exists(),
    reason="outputs/sports_strategy.md not present",
)
def test_real_report_converts_without_error(tmp_path):
    out = tmp_path / "sports_strategy.docx"
    convert(_REAL_REPORT, out)
    assert out.exists() and out.stat().st_size > 1000


@pytest.mark.skipif(
    not _REAL_REPORT.exists(),
    reason="outputs/sports_strategy.md not present",
)
def test_real_report_has_tables(tmp_path):
    out = tmp_path / "sports_strategy.docx"
    convert(_REAL_REPORT, out)
    from docx import Document
    doc = Document(str(out))
    assert len(doc.tables) >= 3
