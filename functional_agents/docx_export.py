"""Markdown → DOCX export for Agent Harness reports.

Single public entry point::

    from functional_agents.docx_export import convert
    convert(Path("outputs/sports_strategy.md"), Path("outputs/sports_strategy.docx"))

Handles: headings (H1–H4), blockquotes/callouts, pipe tables, unordered and
numbered lists, horizontal rules, and inline bold/italic.  All styling is
applied at the run level so the output has no dependency on an external Word
template.
"""
from __future__ import annotations

import re
from pathlib import Path

_FONT = "Calibri"
_INLINE = re.compile(r"(\*\*[^*]+\*\*|\*[^*]+\*)")
_TABLE_SEP = re.compile(r"^\|[-| :]+\|?$")
_ORDERED = re.compile(r"^\d+\. ")


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------

def convert(md_path: Path, docx_path: Path) -> None:
    """Convert *md_path* to *docx_path*.

    Creates parent directories if needed; overwrites any existing DOCX.
    Raises ``ImportError`` if *python-docx* is not installed.
    """
    from docx import Document  # lazy import — optional dependency

    text = md_path.read_text(encoding="utf-8")
    doc = Document()
    _configure_document(doc)
    _render(doc, text)
    docx_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(docx_path))


# ---------------------------------------------------------------------------
# Document configuration
# ---------------------------------------------------------------------------

def _configure_document(doc) -> None:
    from docx.shared import Pt, Inches

    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.0)

    normal = doc.styles["Normal"]
    normal.font.name = _FONT
    normal.font.size = Pt(11)


# ---------------------------------------------------------------------------
# Main renderer (state machine over lines)
# ---------------------------------------------------------------------------

def _render(doc, text: str) -> None:
    lines = text.splitlines()
    i, n = 0, len(lines)

    while i < n:
        line = lines[i]

        if not line.strip():
            i += 1
            continue

        if line.strip() == "---":
            _add_rule(doc)
            i += 1
            continue

        # Headings — longest prefix first
        if line.startswith("#### "):
            _add_heading(doc, line[5:].strip(), level=4)
            i += 1
            continue
        if line.startswith("### "):
            _add_heading(doc, line[4:].strip(), level=3)
            i += 1
            continue
        if line.startswith("## "):
            body = line[3:].strip()
            _add_heading(doc, body, level=2, page_break=_page_break_before(body))
            i += 1
            continue
        if line.startswith("# "):
            _add_heading(doc, line[2:].strip(), level=1)
            i += 1
            continue

        # Blockquote block
        if line.startswith(">"):
            bq: list[str | None] = []
            while i < n:
                l = lines[i]
                if l.startswith(">"):
                    # bare ">" is a paragraph separator, same as a blank line between blocks
                    bq.append(None if l == ">" else l[1:].lstrip())
                    i += 1
                elif l == "" and i + 1 < n and lines[i + 1].startswith(">"):
                    bq.append(None)  # inter-paragraph blank inside block quote
                    i += 1
                else:
                    break
            _add_blockquote(doc, bq)
            continue

        # Table block
        if line.startswith("|"):
            table_lines: list[str] = []
            while i < n and lines[i].startswith("|"):
                table_lines.append(lines[i])
                i += 1
            _add_table(doc, table_lines)
            continue

        # Unordered list
        if line.startswith("- ") or line.startswith("* "):
            while i < n and (lines[i].startswith("- ") or lines[i].startswith("* ")):
                _add_list_item(doc, lines[i][2:], ordered=False)
                i += 1
            continue

        # Ordered list
        if _ORDERED.match(line):
            while i < n and _ORDERED.match(lines[i]):
                _add_list_item(doc, _ORDERED.sub("", lines[i]), ordered=True)
                i += 1
            continue

        # Regular paragraph — collect soft-wrapped continuation lines
        parts = [line]
        i += 1
        while i < n:
            nxt = lines[i]
            if not nxt.strip():
                break
            if nxt.startswith(("#", ">", "|", "- ", "* ")):
                break
            if _ORDERED.match(nxt) or nxt.strip() == "---":
                break
            parts.append(nxt)
            i += 1
        _add_paragraph(doc, " ".join(parts))


def _page_break_before(heading_text: str) -> bool:
    """True when this H2 heading should open on a fresh page."""
    return heading_text.startswith("10.") or heading_text.startswith("Appendix")


# ---------------------------------------------------------------------------
# Block renderers
# ---------------------------------------------------------------------------

def _add_heading(doc, text: str, level: int, *, page_break: bool = False) -> None:
    from docx.shared import Pt, RGBColor

    word_level = {1: 0, 2: 1, 3: 2, 4: 3}.get(level, 1)
    p = doc.add_heading(level=word_level)
    p.clear()

    if page_break:
        p.paragraph_format.page_break_before = True

    _add_runs(p, text)

    sizes    = {1: 20, 2: 14, 3: 12, 4: 11}
    colors   = {1: (0x1F, 0x39, 0x64),
                2: (0x1F, 0x39, 0x64),
                3: (0x2E, 0x4A, 0x7A),
                4: (0x40, 0x40, 0x40)}
    sz       = sizes.get(level, 11)
    rgb      = colors.get(level, (0x20, 0x20, 0x20))

    for run in p.runs:
        run.font.name  = _FONT
        run.font.size  = Pt(sz)
        run.font.color.rgb = RGBColor(*rgb)
        run.font.bold  = level <= 3
        run.font.italic = (level == 4)

    spacing = {2: (Pt(16), Pt(4)), 3: (Pt(10), Pt(2)), 4: (Pt(6), Pt(2))}
    if level in spacing:
        before, after = spacing[level]
        p.paragraph_format.space_before = before
        p.paragraph_format.space_after  = after


def _add_blockquote(doc, segments: list[str | None]) -> None:
    from docx.shared import Pt, RGBColor, Inches
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    # Group on None separators → each group becomes one paragraph
    groups: list[list[str]] = []
    current: list[str] = []
    for item in segments:
        if item is None:
            if current:
                groups.append(current)
            current = []
        else:
            current.append(item)
    if current:
        groups.append(current)

    for group in groups:
        joined = " ".join(group).strip()
        if not joined:
            continue

        p = doc.add_paragraph()
        p.paragraph_format.left_indent  = Inches(0.4)
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after  = Pt(3)

        _add_runs(p, joined)
        for run in p.runs:
            run.font.name = _FONT
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x30, 0x30, 0x30)

        # Blue-grey left border
        pPr  = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        left = OxmlElement("w:left")
        left.set(qn("w:val"),   "single")
        left.set(qn("w:sz"),    "18")
        left.set(qn("w:space"), "4")
        left.set(qn("w:color"), "4472C4")
        pBdr.append(left)
        pPr.append(pBdr)


def _add_table(doc, raw_lines: list[str]) -> None:
    from docx.shared import Pt, RGBColor
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    rows = [
        [c.strip() for c in line.strip().strip("|").split("|")]
        for line in raw_lines
        if not _TABLE_SEP.match(line.strip())
    ]
    if not rows:
        return

    col_count = max(len(r) for r in rows)
    rows = [r + [""] * (col_count - len(r)) for r in rows]

    table = doc.add_table(rows=len(rows), cols=col_count)
    table.style = "Table Grid"

    for r_idx, row_data in enumerate(rows):
        is_header = (r_idx == 0)
        for c_idx, cell_text in enumerate(row_data):
            cell = table.rows[r_idx].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            _add_runs(p, cell_text)
            for run in p.runs:
                run.font.name = _FONT
                run.font.size = Pt(9)
                run.font.bold = is_header
                if is_header:
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            if is_header:
                tc   = cell._tc
                tcPr = tc.get_or_add_tcPr()
                shd  = OxmlElement("w:shd")
                shd.set(qn("w:val"),   "clear")
                shd.set(qn("w:color"), "auto")
                shd.set(qn("w:fill"),  "1F3964")
                tcPr.append(shd)

    doc.add_paragraph()  # breathing room after table


def _add_list_item(doc, text: str, *, ordered: bool) -> None:
    from docx.shared import Pt

    style = "List Number" if ordered else "List Bullet"
    try:
        p = doc.add_paragraph(style=style)
    except KeyError:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Pt(18)
    p.clear()
    _add_runs(p, text)
    for run in p.runs:
        run.font.name = _FONT
        run.font.size = Pt(10)


def _add_paragraph(doc, text: str) -> None:
    from docx.shared import Pt

    p = doc.add_paragraph()
    _add_runs(p, text)
    for run in p.runs:
        run.font.name = _FONT
        run.font.size = Pt(11)
    p.paragraph_format.space_after = Pt(6)


def _add_rule(doc) -> None:
    from docx.shared import Pt
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    p    = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    "4")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), "BBBBBB")
    pBdr.append(bot)
    pPr.append(pBdr)


# ---------------------------------------------------------------------------
# Inline formatting
# ---------------------------------------------------------------------------

def _add_runs(paragraph, text: str) -> None:
    """Add runs to *paragraph* with **bold** and *italic* markers resolved."""
    for part in _INLINE.split(text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**") and len(part) > 4:
            paragraph.add_run(part[2:-2]).bold = True
        elif part.startswith("*") and part.endswith("*") and len(part) > 2:
            paragraph.add_run(part[1:-1]).italic = True
        else:
            paragraph.add_run(part)
